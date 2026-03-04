#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════════════════════╗
║          VAPT ANALYTICAL PLATFORM  — Advanced Edition v2.0                  ║
║  Vulnerability Assessment + Penetration Testing + OSINT Intelligence        ║
║                                                                              ║
║  Inspired by: Vulners, Shodan-style recon, Spiderfoot, SpyFu, Censys,       ║
║               BuiltWith, SecurityHeaders.io, HaveIBeenPwned, DNSDumpster    ║
╚══════════════════════════════════════════════════════════════════════════════╝

Usage:
    python vapt_scanner.py <target_url>
    python vapt_scanner.py https://example.com

Install dependencies:
    pip install requests beautifulsoup4 dnspython matplotlib python-docx
                Pillow tqdm colorama pyOpenSSL cryptography python-whois
"""

import sys, os, re, ssl, json, socket, hashlib, datetime, urllib.parse
import ipaddress, time, threading, warnings, struct, base64
from concurrent.futures import ThreadPoolExecutor, as_completed
warnings.filterwarnings("ignore")

# ── Third-party ───────────────────────────────────────────────────────────────
try:
    import requests
    from requests.packages.urllib3.exceptions import InsecureRequestWarning
    requests.packages.urllib3.disable_warnings(InsecureRequestWarning)
    from bs4 import BeautifulSoup
    import dns.resolver, dns.reversename, dns.zone, dns.query
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.ticker as ticker
    import numpy as np
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from tqdm import tqdm
    from colorama import Fore, Style, init as colorama_init
    colorama_init(autoreset=True)
    from PIL import Image
    import OpenSSL.crypto as crypto
except ImportError as e:
    print(f"\n[!] Missing dependency: {e}")
    print("    pip install requests beautifulsoup4 dnspython matplotlib "
          "python-docx Pillow tqdm colorama pyOpenSSL cryptography")
    sys.exit(1)

try:
    import whois as whois_lib
    WHOIS_AVAILABLE = True
except ImportError:
    WHOIS_AVAILABLE = False

# ─────────────────────────────────────────────────────────────────────────────
#  CONSTANTS & MAPPINGS
# ─────────────────────────────────────────────────────────────────────────────

SEVERITY_ORDER  = ["Critical", "High", "Medium", "Low", "Info"]
SEVERITY_COLORS = {"Critical":"#C00000","High":"#FF4500","Medium":"#FFA500","Low":"#FFD700","Info":"#00BFFF"}
SEVERITY_FILL   = {"Critical":"C00000","High":"FF4500","Medium":"FFA500","Low":"FFD700","Info":"00BFFF"}
SEVERITY_TEXT   = {"Critical":"FFFFFF","High":"FFFFFF","Medium":"000000","Low":"000000","Info":"000000"}

OWASP_MAP = {
    "Missing X-Frame-Options":           ("A05","Security Misconfiguration"),
    "Cross-Frame Scripting (XFS)":       ("A05","Security Misconfiguration"),
    "Missing HSTS Header":               ("A02","Cryptographic Failures"),
    "Missing Content-Security-Policy":   ("A05","Security Misconfiguration"),
    "Missing X-Content-Type-Options":    ("A05","Security Misconfiguration"),
    "Missing Referrer-Policy":           ("A05","Security Misconfiguration"),
    "Missing Permissions-Policy":        ("A05","Security Misconfiguration"),
    "Overly Permissive CORS":            ("A05","Security Misconfiguration"),
    "Server Version Disclosure":         ("A05","Security Misconfiguration"),
    "Open Redirect":                     ("A01","Broken Access Control"),
    "Clickjacking":                      ("A05","Security Misconfiguration"),
    "SSL/TLS Issues":                    ("A02","Cryptographic Failures"),
    "Insecure Cookies":                  ("A02","Cryptographic Failures"),
    "Missing HttpOnly Cookie Flag":      ("A07","Identification and Authentication Failures"),
    "Missing Secure Cookie Flag":        ("A02","Cryptographic Failures"),
    "Cookie Without SameSite":           ("A01","Broken Access Control"),
    "No Input Validation (Reflected)":   ("A03","Injection"),
    "SQL Injection Indicators":          ("A03","Injection"),
    "Directory Listing Enabled":         ("A05","Security Misconfiguration"),
    "DMARC Record Missing":              ("A05","Security Misconfiguration"),
    "DNSSEC Not Implemented":            ("A05","Security Misconfiguration"),
    "SPF Record Missing":                ("A05","Security Misconfiguration"),
    "Outdated Libraries Detected":       ("A06","Vulnerable and Outdated Components"),
    "HTTP Methods Enabled":              ("A05","Security Misconfiguration"),
    "robots.txt Sensitive Disclosure":   ("A01","Broken Access Control"),
    "Exposed Sensitive Files":           ("A05","Security Misconfiguration"),
    "Subresource Integrity Missing":     ("A08","Software and Data Integrity Failures"),
    "Error Message Information Leakage": ("A05","Security Misconfiguration"),
    "Admin Panel Exposed":               ("A01","Broken Access Control"),
    "Backup Files Accessible":           ("A05","Security Misconfiguration"),
    "API Endpoint Exposure":             ("A01","Broken Access Control"),
    "JWT Misconfiguration":              ("A02","Cryptographic Failures"),
    "Cache Control Issues":              ("A02","Cryptographic Failures"),
    "Mixed Content (HTTP in HTTPS)":     ("A02","Cryptographic Failures"),
    "Weak TLS Cipher Suite":             ("A02","Cryptographic Failures"),
    "Certificate Transparency Issues":   ("A02","Cryptographic Failures"),
    "DNS Zone Transfer Possible":        ("A05","Security Misconfiguration"),
    "Subdomain Takeover Risk":           ("A05","Security Misconfiguration"),
    "Open Ports Detected":               ("A05","Security Misconfiguration"),
    "Default Credentials Hint":          ("A07","Identification and Authentication Failures"),
    "Email Addresses Harvested":         ("A01","Broken Access Control"),
    "Sensitive Comments in Source":      ("A05","Security Misconfiguration"),
    "WebSocket Without TLS":             ("A02","Cryptographic Failures"),
    "Path Traversal Indicators":         ("A01","Broken Access Control"),
    "SSRF Surface Detected":             ("A10","Server-Side Request Forgery"),
    "Weak Password Policy Page":         ("A07","Identification and Authentication Failures"),
    "Insecure Deserialization Hints":    ("A08","Software and Data Integrity Failures"),
}

REMEDIATION = {
    "Missing X-Frame-Options":("Add `X-Frame-Options: DENY` or `SAMEORIGIN` to all responses.","Medium"),
    "Missing HSTS Header":("Add `Strict-Transport-Security: max-age=31536000; includeSubDomains; preload`.","Medium"),
    "Missing Content-Security-Policy":("Define a strict CSP header to mitigate XSS.","Medium"),
    "Missing X-Content-Type-Options":("Add `X-Content-Type-Options: nosniff` header.","Low"),
    "Missing Referrer-Policy":("Set `Referrer-Policy: strict-origin-when-cross-origin`.","Low"),
    "Missing Permissions-Policy":("Add `Permissions-Policy` to restrict browser feature access.","Low"),
    "Overly Permissive CORS":("Replace wildcard CORS with a strict trusted-origin allowlist.","Medium"),
    "Server Version Disclosure":("Suppress version info from Server and X-Powered-By headers.","Low"),
    "SSL/TLS Issues":("Use TLS 1.2+ only; renew certificates; enforce HTTPS redirect.","High"),
    "Insecure Cookies":("Set Secure, HttpOnly, and SameSite=Strict on all session cookies.","Medium"),
    "Missing HttpOnly Cookie Flag":("Add HttpOnly flag to all cookies not requiring JS access.","Medium"),
    "Missing Secure Cookie Flag":("Add Secure flag so cookies are only transmitted over HTTPS.","Medium"),
    "Cookie Without SameSite":("Add SameSite=Strict or Lax to all cookies to prevent CSRF.","Medium"),
    "No Input Validation (Reflected)":("Apply server-side input validation and HTML-encode all output.","High"),
    "SQL Injection Indicators":("Use parameterised queries / prepared statements exclusively.","Critical"),
    "Directory Listing Enabled":("Disable directory listing in server config.","Medium"),
    "DMARC Record Missing":("Publish a DMARC record to protect against email spoofing.","Low"),
    "DNSSEC Not Implemented":("Enable DNSSEC at your DNS registrar.","Low"),
    "SPF Record Missing":("Publish an SPF TXT record authorising your mail senders.","Low"),
    "Outdated Libraries Detected":("Update all third-party JS libraries to latest stable versions.","Medium"),
    "HTTP Methods Enabled":("Disable TRACE, TRACK, PUT, DELETE unless explicitly required.","Medium"),
    "robots.txt Sensitive Disclosure":("Remove sensitive paths from robots.txt; use real access controls.","Low"),
    "Exposed Sensitive Files":("Block access to .env, .git, config files via web server deny rules.","High"),
    "Subresource Integrity Missing":("Add integrity and crossorigin attributes to all external resources.","Low"),
    "Error Message Information Leakage":("Return generic error messages; log details server-side only.","Medium"),
    "Admin Panel Exposed":("Move admin interfaces behind VPN or IP whitelist; enforce MFA.","High"),
    "Backup Files Accessible":("Remove all backup files from web root; deny .bak/.sql in server config.","High"),
    "API Endpoint Exposure":("Require authentication on all API endpoints; apply rate limiting.","High"),
    "JWT Misconfiguration":("Validate JWT algorithm, expiry, and signature; reject 'none' algorithm.","High"),
    "Cache Control Issues":("Add `Cache-Control: no-store` on sensitive pages.","Medium"),
    "Mixed Content (HTTP in HTTPS)":("Load all resources over HTTPS; update all hardcoded HTTP URLs.","Medium"),
    "Weak TLS Cipher Suite":("Disable weak ciphers (RC4, DES, 3DES); use ECDHE+AES256.","High"),
    "Certificate Transparency Issues":("Monitor CT logs for unauthorised certificate issuance.","Medium"),
    "DNS Zone Transfer Possible":("Restrict AXFR transfers to authorised secondary DNS IPs only.","High"),
    "Subdomain Takeover Risk":("Remove or re-point dangling DNS records to active services.","High"),
    "Open Ports Detected":("Close or firewall unnecessary open ports; document all services.","Medium"),
    "Default Credentials Hint":("Change default credentials immediately; enforce strong password policy.","Critical"),
    "Email Addresses Harvested":("Obfuscate email addresses on public pages to prevent harvesting.","Low"),
    "Sensitive Comments in Source":("Remove developer comments with paths, credentials, or logic from production.","Low"),
    "WebSocket Without TLS":("Use wss:// (WebSocket Secure) instead of ws://.","Medium"),
    "Path Traversal Indicators":("Validate and sanitise all file path inputs; use allowlists.","High"),
    "SSRF Surface Detected":("Validate and whitelist all server-side URL fetch targets; block internal IPs.","High"),
    "Open Redirect":("Validate redirect URLs against a strict whitelist server-side.","Medium"),
    "Clickjacking":("Set X-Frame-Options: DENY and CSP frame-ancestors 'none'.","Medium"),
    "Weak Password Policy Page":("Enforce 12+ char passwords, rate limiting, and MFA.","High"),
    "Insecure Deserialization Hints":("Avoid deserialising untrusted data; use safe serialisation formats.","High"),
    "Cross-Frame Scripting (XFS)":("Implement X-Frame-Options: SAMEORIGIN and CSP frame-ancestors 'self'.","Medium"),
}


SESSION = requests.Session()
SESSION.headers.update({"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"})
SESSION.verify = False


# ─────────────────────────────────────────────────────────────────────────────
#  HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def banner():
    print(Fore.CYAN + """
╔══════════════════════════════════════════════════════════════════════╗
║  ██╗   ██╗ █████╗ ██████╗ ████████╗    ██████╗ ██████╗  ██████╗    ║
║  ██║   ██║██╔══██╗██╔══██╗╚══██╔══╝    ██╔══██╗██╔══██╗██╔═══██╗   ║
║  ██║   ██║███████║██████╔╝   ██║       ██████╔╝██████╔╝██║   ██║   ║
║  ╚██╗ ██╔╝██╔══██║██╔═══╝    ██║       ██╔═══╝ ██╔══██╗██║   ██║   ║
║   ╚████╔╝ ██║  ██║██║        ██║       ██║     ██║  ██║╚██████╔╝   ║
║    ╚═══╝  ╚═╝  ╚═╝╚═╝        ╚═╝       ╚═╝     ╚═╝  ╚═╝ ╚═════╝    ║
║                                                                      ║
║    VAPT Analytical Platform v2.0  ·  Advanced OSINT + Security       ║
╚══════════════════════════════════════════════════════════════════════╝""" + Style.RESET_ALL)

def normalise_url(url):
    if not url.startswith(("http://","https://")):
        url = "https://" + url
    return url.rstrip("/")

def get_domain(url):
    return urllib.parse.urlparse(url).netloc.split(":")[0]

def safe_get(url, timeout=10, allow_redirects=True, headers=None):
    try:
        h = dict(SESSION.headers)
        if headers:
            h.update(headers)
        return SESSION.get(url, timeout=timeout, allow_redirects=allow_redirects, headers=h, verify=False)
    except Exception:
        return None

def hex_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2],16) for i in (0,2,4))

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def ip_to_geo(ip):
    try:
        r = requests.get(f"http://ip-api.com/json/{ip}?fields=country,regionName,city,isp,org,as,query",
                         timeout=6)
        if r.status_code == 200:
            return r.json()
    except Exception:
        pass
    return {}


# ─────────────────────────────────────────────────────────────────────────────
#  MAIN SCANNER CLASS
# ─────────────────────────────────────────────────────────────────────────────

class VAPTScanner:
    def __init__(self, target):
        self.target   = normalise_url(target)
        self.domain   = get_domain(self.target)
        self.findings = []
        self.intel    = {}
        self._resp    = None

    def _add(self, name, evidence, detail, urls=None, severity=None):
        if name not in OWASP_MAP or name not in REMEDIATION:
            return
        owasp_id, owasp_name = OWASP_MAP.get(name,("A05","Security Misconfiguration"))
        remedy, sev = REMEDIATION.get(name,("Review and remediate.","Info"))
        if severity:
            sev = severity
        self.findings.append({
            "id":        f"V-{len(self.findings)+1:02d}",
            "name":      name,
            "severity":  sev,
            "owasp_id":  owasp_id,
            "owasp_name":owasp_name,
            "evidence":  evidence,
            "detail":    detail,
            "remedy":    remedy,
            "urls":      urls or [self.target],
        })

    def _main(self):
        if self._resp is None:
            self._resp = safe_get(self.target)
        return self._resp

    # ══════════════════════════════════════════════════════════════════════
    #   OSINT & RECON MODULES
    # ══════════════════════════════════════════════════════════════════════

    def recon_dns(self):
        info = {}
        for rtype in ["A","AAAA","MX","NS","TXT","CNAME","SOA","CAA"]:
            try:
                answers = dns.resolver.resolve(self.domain, rtype, lifetime=5)
                info[rtype] = [str(r) for r in answers]
            except Exception:
                info[rtype] = []
        ptr_map = {}
        for ip in info.get("A",[])[:3]:
            try:
                rev = dns.reversename.from_address(ip)
                ptr_answers = dns.resolver.resolve(rev,"PTR",lifetime=4)
                ptr_map[ip] = str(ptr_answers[0])
            except Exception:
                ptr_map[ip] = "No PTR record"
        info["PTR_MAP"] = ptr_map
        self.intel["dns"] = info

    def recon_ip_info(self):
        ips = []
        try:
            ips = [str(r) for r in dns.resolver.resolve(self.domain,"A",lifetime=5)]
        except Exception:
            try:
                ips = [socket.gethostbyname(self.domain)]
            except Exception:
                pass
        geo_results = []
        for ip in ips[:3]:
            geo = ip_to_geo(ip)
            geo_results.append({"ip":ip, **geo})
        self.intel["ip_info"] = geo_results

    def recon_whois(self):
        data = {}
        if WHOIS_AVAILABLE:
            try:
                w = whois_lib.whois(self.domain)
                def first(v):
                    return str(v[0] if isinstance(v,list) else v or "N/A")
                data = {
                    "registrar":    first(w.registrar),
                    "creation_date":first(w.creation_date),
                    "expiry_date":  first(w.expiration_date),
                    "updated_date": first(w.updated_date),
                    "name_servers": list(set([str(n).lower() for n in (w.name_servers or [])])),
                    "status":       first(w.status),
                    "org":          str(w.org or w.registrant_name or "N/A"),
                    "country":      str(w.country or "N/A"),
                    "emails":       list(set(w.emails)) if isinstance(w.emails,list) else ([w.emails] if w.emails else []),
                }
                exp = w.expiration_date
                if isinstance(exp,list): exp = exp[0]
                if exp and isinstance(exp,datetime.datetime):
                    days = (exp - datetime.datetime.utcnow()).days
                    if days < 30:
                        self._add("SSL/TLS Issues",
                            f"Domain expires in {days} days ({exp.date()})",
                            "Domain registration nearing expiry — risk of accidental lapse.",
                            severity="High")
            except Exception as e:
                data = {"error": str(e)}
        else:
            try:
                r = requests.get(f"https://rdap.verisign.com/com/v1/domain/{self.domain}",timeout=8)
                if r.status_code == 200:
                    j = r.json()
                    for event in j.get("events",[]):
                        if event.get("eventAction") == "registration":
                            data["creation_date"] = event.get("eventDate","N/A")
                        if event.get("eventAction") == "expiration":
                            data["expiry_date"] = event.get("eventDate","N/A")
            except Exception:
                data["note"] = "Install python-whois for full WHOIS data: pip install python-whois"
        self.intel["whois"] = data

    def recon_ssl_certificate(self):
        cert_info = {}
        try:
            ctx = ssl.create_default_context()
            ctx.check_hostname = False
            ctx.verify_mode = ssl.CERT_NONE
            with socket.create_connection((self.domain,443),timeout=8) as sock:
                with ctx.wrap_socket(sock, server_hostname=self.domain) as ssock:
                    raw = ssock.getpeercert(binary_form=True)
                    pem = ssl.DER_cert_to_PEM_cert(raw)
                    cert_info["protocol"] = ssock.version()
                    cert_info["cipher"]   = ssock.cipher()
                    x509 = crypto.load_certificate(crypto.FILETYPE_PEM, pem)
                    subj = x509.get_subject()
                    cert_info["subject"]  = {"CN":subj.CN,"O":subj.O,"C":subj.C}
                    issuer = x509.get_issuer()
                    cert_info["issuer"]   = {"CN":issuer.CN,"O":issuer.O}
                    cert_info["serial"]   = str(x509.get_serial_number())
                    cert_info["not_after"]= x509.get_notAfter().decode()
                    cert_info["sha256_fp"]= x509.digest("sha256").decode()
                    sans = []
                    for i in range(x509.get_extension_count()):
                        ext = x509.get_extension(i)
                        if ext.get_short_name() == b"subjectAltName":
                            sans = [s.strip() for s in str(ext).split(",")]
                    cert_info["san"] = sans
                    exp_str = x509.get_notAfter().decode()
                    exp_dt = datetime.datetime.strptime(exp_str,"%Y%m%d%H%M%SZ")
                    days = (exp_dt - datetime.datetime.utcnow()).days
                    cert_info["days_until_expiry"] = days
                    if days < 0:
                        self._add("SSL/TLS Issues",f"Certificate EXPIRED {abs(days)} days ago",
                            "SSL certificate has expired — all HTTPS connections show a security warning.",severity="Critical")
                    elif days < 30:
                        self._add("SSL/TLS Issues",f"Certificate expires in {days} days",
                            "Certificate near expiry — renew immediately.",severity="High")
                    proto = ssock.version()
                    if proto in ("SSLv2","SSLv3","TLSv1","TLSv1.1"):
                        self._add("Weak TLS Cipher Suite",f"Protocol: {proto}",
                            f"{proto} is deprecated and vulnerable to known attacks (BEAST, POODLE).",severity="High")
                    cipher_name = ssock.cipher()[0] if ssock.cipher() else ""
                    for wc in ["RC4","DES","3DES","EXPORT","NULL","anon"]:
                        if wc.lower() in cipher_name.lower():
                            self._add("Weak TLS Cipher Suite",f"Weak cipher: {cipher_name}",
                                f"Cipher {cipher_name} is cryptographically weak.",severity="High")
                            break
        except ssl.SSLError as e:
            cert_info["error"] = str(e)
            self._add("SSL/TLS Issues",f"SSL error: {e}","SSL/TLS configuration error.",severity="High")
        except Exception as e:
            cert_info["error"] = str(e)
        self.intel["ssl"] = cert_info

    def recon_subdomains(self):
        wordlist = [
            "www","mail","remote","blog","webmail","server","ns1","ns2","smtp","secure",
            "vpn","m","shop","ftp","mail2","test","portal","ns","ww1","host","support",
            "dev","web","bbs","ww42","mx","email","1","mail1","2","forum","owa","www2",
            "gw","admin","store","mx1","cdn","api","exchange","app","archive","beta",
            "cpanel","whm","autodiscover","autoconfig","irc","news","media","crm",
            "staging","office","chat","direct","dashboard","login","auth","status",
            "monitor","db","mysql","oracle","sql","demo","git","svn","jira","confluence",
            "jenkins","grafana","kibana","docs","help","wiki","intranet","internal",
            "vpn2","remote2","extranet","partner","customer","portal2","new","old",
        ]
        found = []
        takeover_services = ["s3.amazonaws.com","github.io","heroku.com","azurewebsites.net",
                             "cloudfront.net","wordpress.com","shopify.com","fastly.net"]

        def check_sub(sub):
            fqdn = f"{sub}.{self.domain}"
            try:
                answers = dns.resolver.resolve(fqdn,"A",lifetime=2)
                ips = [str(r) for r in answers]
                return {"subdomain":fqdn,"ips":ips}
            except Exception:
                return None

        with ThreadPoolExecutor(max_workers=40) as exe:
            for result in exe.map(check_sub, wordlist):
                if result:
                    found.append(result)

        for sub in found:
            fqdn = sub["subdomain"]
            try:
                cnames = dns.resolver.resolve(fqdn,"CNAME",lifetime=3)
                for cn in cnames:
                    cval = str(cn.target).lower()
                    for svc in takeover_services:
                        if svc in cval:
                            r = safe_get(f"https://{fqdn}",timeout=4)
                            if not r or r.status_code in (404,503):
                                self._add("Subdomain Takeover Risk",
                                    f"{fqdn} CNAME → {cval} (service may be unclaimed)",
                                    "CNAME points to an unclaimed cloud service — attacker can claim it.",
                                    [f"https://{fqdn}"],severity="High")
            except Exception:
                pass

        self.intel["subdomains"] = sorted(found, key=lambda x:x["subdomain"])

    def recon_port_scan(self):
        PORTS = {
            21:"FTP",22:"SSH",23:"Telnet",25:"SMTP",53:"DNS",
            80:"HTTP",110:"POP3",143:"IMAP",443:"HTTPS",445:"SMB",
            3306:"MySQL",3389:"RDP",5432:"PostgreSQL",5900:"VNC",
            6379:"Redis",8080:"HTTP-Alt",8443:"HTTPS-Alt",8888:"HTTP-Dev",
            9200:"Elasticsearch",27017:"MongoDB",11211:"Memcached",
            2375:"Docker API",4848:"GlassFish",7001:"WebLogic",
            8161:"ActiveMQ",9090:"Cockpit",6443:"Kubernetes API",
        }
        try:
            ip = socket.gethostbyname(self.domain)
        except Exception:
            self.intel["ports"] = []
            return

        open_ports = []
        def check_port(port):
            try:
                with socket.create_connection((ip,port),timeout=1.2):
                    return port
            except Exception:
                return None

        with ThreadPoolExecutor(max_workers=60) as exe:
            results = list(exe.map(check_port, PORTS.keys()))

        for port in results:
            if port:
                service = PORTS.get(port,"Unknown")
                open_ports.append({"port":port,"service":service})
                if port in (23,445,3389,5900,6379,9200,27017,11211,2375,4848,7001):
                    self._add("Open Ports Detected",f"Port {port} ({service}) open on {ip}",
                        f"{service} on port {port} should not be publicly accessible.",severity="High")
                elif port in (21,3306,5432,8888):
                    self._add("Open Ports Detected",f"Port {port} ({service}) open",
                        f"{service} port publicly accessible — restrict to internal network.",severity="Medium")

        self.intel["ports"] = open_ports

    def recon_technologies(self):
        r = self._main()
        if not r:
            self.intel["technologies"] = {}
            return
        body    = r.text.lower()
        headers = {k.lower():v for k,v in r.headers.items()}
        soup    = BeautifulSoup(r.text,"html.parser")
        tech    = {}

        cms_sigs = {
            "WordPress":   ["wp-content","wp-includes","wp-json"],
            "Drupal":      ["drupal.js","drupal.min.js","/sites/default/"],
            "Joomla":      ["/components/com_","joomla"],
            "Magento":     ["mage/","magento","varien"],
            "Shopify":     ["cdn.shopify.com","shopify.com/s/"],
            "Wix":         ["wixsite.com","wix.com/"],
            "Squarespace": ["squarespace.com","static1.squarespace"],
            "Ghost":       ["ghost.io","content/images/"],
            "PrestaShop":  ["prestashop","presta_shop"],
            "OpenCart":    ["opencart","route=common"],
        }
        for name, sigs in cms_sigs.items():
            if any(s in body for s in sigs):
                tech["CMS"] = name
                break

        server = headers.get("server","")
        if server: tech["Server"] = server
        powered = headers.get("x-powered-by","")
        if powered: tech["Framework"] = powered

        cdn_sigs = {
            "Cloudflare":        ["cf-ray","cf-cache-status"],
            "Fastly":            ["x-served-by","x-cache"],
            "Akamai":            ["x-check-cacheable","x-akamai"],
            "Amazon CloudFront": ["x-amz-cf-id","cloudfront"],
            "Azure CDN":         ["x-azure-ref"],
            "BunnyCDN":          ["bunnycdn","bunny.net"],
        }
        for cdn, hdrs in cdn_sigs.items():
            if any(h in headers or h in body for h in hdrs):
                tech["CDN"] = cdn
                break

        waf_sigs = {
            "Cloudflare WAF": ["cf-ray"],
            "AWS WAF":        ["x-amzn-requestid","awselb"],
            "Sucuri":         ["x-sucuri-id"],
            "Incapsula":      ["x-iinfo","incap_ses"],
            "F5 BIG-IP":      ["bigipserver"],
            "Barracuda":      ["barra_counter_session"],
            "ModSecurity":    ["mod_security","modsec"],
            "Imperva":        ["x-iinfo"],
        }
        detected_waf = "None detected"
        for waf, sigs in waf_sigs.items():
            if any(s in headers or s in body for s in sigs):
                detected_waf = waf
                break
        tech["WAF"] = detected_waf

        js_sigs = {
            "React":     ["react.js","react.min.js","react-dom","__reactfiber"],
            "Vue.js":    ["vue.js","vue.min.js","__vue__"],
            "Angular":   ["angular.js","ng-version","angular.min"],
            "jQuery":    ["jquery.js","jquery.min.js","jquery-"],
            "Bootstrap": ["bootstrap.js","bootstrap.min","bootstrap.css"],
            "Next.js":   ["_next/","__next"],
            "Nuxt.js":   ["_nuxt/","nuxtjs"],
            "Ember.js":  ["ember.js","ember.min"],
            "Svelte":    ["svelte","__svelte"],
            "Alpine.js": ["alpine.js","x-data="],
        }
        frameworks = [name for name,sigs in js_sigs.items() if any(s in body for s in sigs)]
        if frameworks: tech["JS Frameworks"] = ", ".join(frameworks)

        analytics_sigs = {
            "Google Analytics":   ["google-analytics.com","gtag/js","UA-","G-"],
            "Google Tag Manager": ["googletagmanager.com"],
            "Facebook Pixel":     ["connect.facebook.net","fbq("],
            "HotJar":             ["hotjar.com","hjSetting"],
            "Mixpanel":           ["mixpanel.com"],
            "Segment.io":         ["segment.com","analytics.js"],
            "Intercom":           ["intercom.io"],
            "Crisp":              ["crisp.chat"],
            "Zendesk":            ["zendesk.com","zopim"],
        }
        analytics = [name for name,sigs in analytics_sigs.items() if any(s in body for s in sigs)]
        tech["Analytics/Tracking"] = analytics or ["None detected"]

        mx_records = self.intel.get("dns",{}).get("MX",[])
        if mx_records:
            mx_str = " ".join(mx_records).lower()
            if "google" in mx_str or "gmail" in mx_str:
                tech["Email Provider"] = "Google Workspace"
            elif "outlook" in mx_str or "microsoft" in mx_str:
                tech["Email Provider"] = "Microsoft 365"
            elif "amazonses" in mx_str:
                tech["Email Provider"] = "Amazon SES"
            else:
                tech["Email Provider"] = mx_records[0] if mx_records else "Unknown"

        # Payment processors
        payment_sigs = {"Stripe":["stripe.com","stripe.js"],"PayPal":["paypal.com","paypaljs"],
                        "Braintree":["braintree","braintreepayments"],"Square":["squareup.com"]}
        payments = [n for n,s in payment_sigs.items() if any(x in body for x in s)]
        if payments: tech["Payment Processors"] = ", ".join(payments)

        self.intel["technologies"] = tech

    def recon_email_harvest(self):
        r = self._main()
        if not r: return
        emails = list(set(re.findall(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", r.text)))
        emails = [e for e in emails if not any(x in e.lower() for x in
                  ["example.com","w3.org","schema.org","sentry.io","jquery.com"])]
        if emails:
            self._add("Email Addresses Harvested",
                f"{len(emails)} email(s): {', '.join(emails[:5])}",
                "Email addresses in page source can be harvested for phishing and spam.")
        self.intel["emails"] = emails

    def recon_source_comments(self):
        r = self._main()
        if not r: return
        comments = re.findall(r"<!--(.*?)-->", r.text, re.DOTALL)
        keywords = ["todo","fixme","hack","bug","password","secret","key","token",
                    "admin","credential","database","api","remove","temp","debug","localhost"]
        interesting = []
        for c in comments:
            c_clean = c.strip()
            if any(k in c_clean.lower() for k in keywords) and len(c_clean) > 5:
                interesting.append(c_clean[:200])
        if interesting:
            self._add("Sensitive Comments in Source",
                f"{len(interesting)} sensitive comment(s) found.",
                "HTML comments contain sensitive keywords (passwords, API keys, TODOs).")
        self.intel["source_comments"] = interesting

    def recon_linked_pages(self):
        r = self._main()
        if not r: return
        soup = BeautifulSoup(r.text,"html.parser")
        links = set()
        base = f"https://{self.domain}"
        for a in soup.find_all("a",href=True):
            href = a["href"]
            if href.startswith("/"):
                links.add(base + href)
            elif self.domain in href:
                links.add(href)
        self.intel["internal_links"] = list(links)[:50]

    def recon_ct_logs(self):
        certs = []
        try:
            r = requests.get(f"https://crt.sh/?q=%.{self.domain}&output=json",
                             timeout=10, headers={"User-Agent":"VAPT-Scanner/2.0"})
            if r.status_code == 200:
                data = r.json()
                seen = set()
                for entry in data[:100]:
                    name = entry.get("name_value","").strip()
                    logged = entry.get("entry_timestamp","")
                    for n in name.split("\n"):
                        n = n.strip().lower()
                        if n and n not in seen:
                            seen.add(n)
                            certs.append({"domain":n,"logged":logged[:10] if logged else "unknown"})
        except Exception:
            pass
        self.intel["ct_logs"] = certs[:50]
        wild = [c for c in certs if c["domain"].startswith("*")]
        if wild:
            self._add("Certificate Transparency Issues",
                f"{len(wild)} wildcard certificate(s) in CT logs.",
                "Wildcard certificates in public CT logs — monitor for unexpected issuance.")

    def recon_waf_detection(self):
        """Advanced WAF/firewall detection via probe responses."""
        waf_info = {"detected":False,"product":"None"}
        payloads = ["?q=<script>alert(1)</script>","?id=1' OR '1'='1","?file=../../../../etc/passwd"]
        for p in payloads:
            r = safe_get(self.target + p, timeout=6, allow_redirects=False)
            if r:
                hdrs = {k.lower():v for k,v in r.headers.items()}
                if r.status_code in (403,406,429,503):
                    waf_info["detected"] = True
                    if "cf-ray" in hdrs:
                        waf_info["product"] = "Cloudflare"
                    elif "x-sucuri-id" in hdrs:
                        waf_info["product"] = "Sucuri"
                    elif "x-iinfo" in hdrs:
                        waf_info["product"] = "Incapsula/Imperva"
                    else:
                        waf_info["product"] = "Unknown WAF"
                    break
        self.intel["waf"] = waf_info

    # ══════════════════════════════════════════════════════════════════════
    #   VULNERABILITY CHECKS (40+)
    # ══════════════════════════════════════════════════════════════════════

    def check_security_headers(self):
        r = self._main()
        if not r: return
        headers = {k.lower():v for k,v in r.headers.items()}
        checks = [
            ("x-frame-options","Missing X-Frame-Options","X-Frame-Options absent — pages embeddable in iframes, enabling clickjacking."),
            ("strict-transport-security","Missing HSTS Header","HSTS absent — browsers may connect over plain HTTP, enabling SSL-stripping."),
            ("content-security-policy","Missing Content-Security-Policy","No CSP — XSS and data-injection risk significantly elevated."),
            ("x-content-type-options","Missing X-Content-Type-Options","X-Content-Type-Options absent — MIME-type sniffing possible."),
            ("referrer-policy","Missing Referrer-Policy","Referrer-Policy absent — URL parameters may leak to third-party sites."),
            ("permissions-policy","Missing Permissions-Policy","Permissions-Policy absent — browser features unrestricted for embedded content."),
        ]
        for hdr, vuln, detail in checks:
            if hdr not in headers:
                self._add(vuln, f"Header '{hdr}' not present.", detail)

    def check_cors(self):
        r = self._main()
        if not r: return
        acao = r.headers.get("Access-Control-Allow-Origin","")
        acac = r.headers.get("Access-Control-Allow-Credentials","")
        if acao.strip() == "*":
            detail = "Wildcard CORS — any origin can make cross-origin requests."
            sev = "Critical" if acac.lower() == "true" else "Medium"
            self._add("Overly Permissive CORS",f"ACAO: {acao} ACAC: {acac}",detail,severity=sev)

    def check_cors_preflight(self):
        try:
            r = SESSION.options(self.target,timeout=7,verify=False,headers={
                "Origin":"https://evil-attacker.com",
                "Access-Control-Request-Method":"GET",
            })
            acao = r.headers.get("Access-Control-Allow-Origin","")
            if "evil-attacker.com" in acao:
                self._add("Overly Permissive CORS",
                    f"Server reflects arbitrary Origin: {acao}",
                    "Origin header reflected without validation — full CORS bypass possible.",severity="High")
        except Exception:
            pass

    def check_server_disclosure(self):
        r = self._main()
        if not r: return
        tokens = []
        for h in ["Server","X-Powered-By","X-AspNet-Version","X-AspNetMvc-Version","X-Generator"]:
            v = r.headers.get(h,"")
            if v: tokens.append(f"{h}: {v}")
        if tokens:
            self._add("Server Version Disclosure","; ".join(tokens),
                "Server advertising software/version — aids attacker CVE research.")

    def check_ssl(self):
        http_url = self.target.replace("https://","http://",1)
        r = safe_get(http_url,allow_redirects=False,timeout=6)
        if r and r.status_code not in (301,302,307,308):
            self._add("SSL/TLS Issues",f"HTTP {http_url} returned {r.status_code} (no HTTPS redirect)",
                "Plain HTTP not redirected to HTTPS — data transmitted in cleartext.",severity="High")

    def check_clickjacking(self):
        r = self._main()
        if not r: return
        xfo = r.headers.get("X-Frame-Options","").upper()
        csp = r.headers.get("Content-Security-Policy","")
        if not xfo and "frame-ancestors" not in csp.lower():
            self._add("Clickjacking","Neither X-Frame-Options nor CSP frame-ancestors found.",
                "Page embeddable in any iframe — trivially clickjackable.")

    def check_mixed_content(self):
        r = self._main()
        if not r or not self.target.startswith("https://"): return
        http_links = re.findall(r'(?:src|href)=["\']http://[^"\']+["\']',r.text,re.I)
        if http_links:
            self._add("Mixed Content (HTTP in HTTPS)",
                f"{len(http_links)} HTTP resource(s) in HTTPS page: {http_links[0][:80]}",
                "HTTP resources in HTTPS page — assets exposed to MITM interception.")

    def check_cookies(self):
        r = self._main()
        if not r: return
        reported = set()
        for cookie in r.cookies:
            if not cookie.secure and "Missing Secure Cookie Flag" not in reported:
                reported.add("Missing Secure Cookie Flag")
                self._add("Missing Secure Cookie Flag",f"Cookie '{cookie.name}' missing Secure flag.",
                    "Cookie transmitted over plain HTTP — intercept risk.")
            if not cookie.has_nonstandard_attr("httponly") and "Missing HttpOnly Cookie Flag" not in reported:
                reported.add("Missing HttpOnly Cookie Flag")
                self._add("Missing HttpOnly Cookie Flag",f"Cookie '{cookie.name}' missing HttpOnly flag.",
                    "Cookie accessible to JavaScript — XSS-based theft possible.")
            if not cookie.has_nonstandard_attr("samesite") and "Cookie Without SameSite" not in reported:
                reported.add("Cookie Without SameSite")
                self._add("Cookie Without SameSite",f"Cookie '{cookie.name}' missing SameSite attribute.",
                    "Cookie sent on cross-origin requests — CSRF attack vector.")

    def check_cache_control(self):
        r = self._main()
        if not r: return
        cc = r.headers.get("Cache-Control","").lower()
        body_lower = r.text.lower()
        if "no-store" not in cc and "no-cache" not in cc and "private" not in cc:
            if any(k in body_lower for k in ["password","logout","account","dashboard","session","profile"]):
                self._add("Cache Control Issues",f"Cache-Control: {r.headers.get('Cache-Control','not set')}",
                    "Sensitive page without no-store/no-cache — may be cached, leaking session data.")

    def check_directory_listing(self):
        for path in ["/images/","/uploads/","/assets/","/static/","/files/","/backup/","/data/","/tmp/"]:
            url = self.target + path
            r = safe_get(url,timeout=5)
            if r and r.status_code == 200 and ("index of" in r.text.lower() or "parent directory" in r.text.lower()):
                self._add("Directory Listing Enabled",f"Directory listing at {url}",
                    "Web server lists directory contents — internal file structures exposed.",[url])
                break

    def check_sensitive_files(self):
        targets = [
            ("/.env","Environment file"),("/.env.local","Local env"),("/.env.production","Prod env"),
            ("/.git/config","Git config"),("/.git/HEAD","Git HEAD"),
            ("/config.php","PHP config"),("/wp-config.php","WP config"),
            ("/wp-config.php.bak","WP config backup"),("/.htaccess","Apache config"),
            ("/server-status","Apache status"),("/phpinfo.php","PHP info"),
            ("/info.php","PHP info"),("/adminer.php","Adminer"),("/phpmyadmin/","phpMyAdmin"),
            ("/database.sql","SQL dump"),("/backup.sql","SQL backup"),
            ("/web.config","IIS config"),("/crossdomain.xml","Flash crossdomain"),
            ("/.DS_Store","macOS DS_Store"),("/npm-debug.log","npm log"),
            ("/yarn-error.log","Yarn log"),("/composer.json","Composer manifest"),
            ("/package.json","NPM manifest"),("/.travis.yml","CI config"),
            ("/Dockerfile","Dockerfile"),("/docker-compose.yml","Docker Compose"),
        ]
        for path, label in targets:
            url = self.target + path
            r = safe_get(url,timeout=5)
            if r and r.status_code == 200 and len(r.text.strip()) > 5:
                if any(kw in r.text.lower() for kw in ["db_","database","password","secret","api_key","php","[core]","ref:","from","env","port"]):
                    self._add("Exposed Sensitive Files",f"{label} at {url}",
                        f"'{path}' publicly accessible — may expose credentials or configuration.",[url],severity="High")

    def check_admin_panels(self):
        admin_paths = [
            "/admin","/admin/","/administrator","/wp-admin","/wp-login.php",
            "/login","/panel","/cpanel","/whm","/webmail","/manage","/management",
            "/backend","/cms","/console","/phpmyadmin","/adminer","/jenkins",
            "/grafana","/kibana","/portainer","/traefik","/vault","/rancher",
            "/solr","/mongo-express","/redis-commander","/flower","/airflow",
        ]
        for path in admin_paths:
            url = self.target + path
            r = safe_get(url,timeout=5,allow_redirects=True)
            if r and r.status_code == 200:
                if any(k in r.text.lower() for k in ["login","password","username","sign in","administration"]):
                    self._add("Admin Panel Exposed",f"Admin interface at {url}",
                        "Administrative interface publicly accessible — brute-force/credential stuffing target.",[url],severity="High")
                    break

    def check_backup_files(self):
        base = self.domain.replace(".","_").replace("-","_")
        today = datetime.date.today()
        patterns = []
        for ext in [".bak",".zip",".tar.gz",".tar",".gz",".old",".sql",".dump"]:
            patterns.extend([f"/{base}{ext}",f"/backup{ext}",f"/site{ext}",
                             f"/backup_{today.year}{today.month:02d}{ext}"])
        for path in patterns[:12]:
            url = self.target + path
            r = safe_get(url,timeout=5)
            if r and r.status_code == 200 and len(r.content) > 100:
                ct = r.headers.get("Content-Type","")
                if any(t in ct for t in ["zip","gzip","octet-stream","sql","tar"]):
                    self._add("Backup Files Accessible",f"Backup at {url} ({len(r.content)} bytes)",
                        "Archive/backup file publicly downloadable — may contain source code, credentials.",[url],severity="High")
                    break

    def check_api_exposure(self):
        api_paths = [
            "/api","/api/v1","/api/v2","/api/v3","/graphql","/rest",
            "/swagger","/swagger-ui.html","/api-docs","/openapi.json",
            "/swagger.json","/swagger.yaml","/.well-known/api-catalog",
            "/api/users","/api/user","/api/admin","/api/keys","/api/config",
        ]
        for path in api_paths:
            url = self.target + path
            r = safe_get(url,timeout=6)
            if r and r.status_code == 200:
                ct = r.headers.get("Content-Type","").lower()
                if "json" in ct or "yaml" in ct or "swagger" in r.text.lower() or "graphql" in r.text.lower():
                    self._add("API Endpoint Exposure",f"Unauthenticated API at {url}",
                        "API endpoint responds without authentication.",[url])
                    break

    def check_http_methods(self):
        try:
            r = SESSION.options(self.target,timeout=7,verify=False)
            allow = r.headers.get("Allow","")
            dangerous = [m for m in ["TRACE","TRACK","PUT","DELETE"] if m in allow.upper()]
            if dangerous:
                self._add("HTTP Methods Enabled",f"Allow: {allow}",
                    f"Dangerous HTTP methods: {', '.join(dangerous)} — may allow server manipulation.")
        except Exception:
            pass

    def check_robots(self):
        r = safe_get(self.target+"/robots.txt",timeout=6)
        if r and r.status_code == 200:
            sensitive = re.findall(r"(?i)(?:disallow|allow):\s*(/[^\n]*(?:admin|backup|config|private|secret|api|db|database|internal|staging)[^\n]*)",r.text)
            if sensitive:
                self._add("robots.txt Sensitive Disclosure",
                    f"Sensitive paths: {', '.join(sensitive[:5])}",
                    "robots.txt discloses internal paths — guides attackers to high-value targets.")

    def check_js_libraries(self):
        r = self._main()
        if not r: return
        soup = BeautifulSoup(r.text,"html.parser")
        VULN_LIBS = {
            "jquery-1.":  "jQuery 1.x — multiple XSS CVEs (CVE-2019-11358, CVE-2020-11022/23)",
            "jquery-2.":  "jQuery 2.x — EOL, CVE-2020-11022",
            "jquery-3.4": "jQuery 3.4.x — prototype pollution CVE-2019-11358",
            "angular.js": "AngularJS 1.x — EOL Dec 2021, multiple XSS vectors",
            "bootstrap-2":"Bootstrap 2.x — EOL, no security patches",
            "bootstrap-3":"Bootstrap 3.x — EOL, no security patches",
            "prototype-1":"Prototype.js — prototype pollution vulnerabilities",
        }
        hits = []
        for tag in soup.find_all("script",src=True):
            src = tag["src"].lower()
            for pattern, note in VULN_LIBS.items():
                if pattern in src:
                    hits.append(note)
        if hits:
            self._add("Outdated Libraries Detected","; ".join(hits[:3]),
                "Vulnerable JS libraries loaded — known CVEs exploitable client-side.")

    def check_sri(self):
        r = self._main()
        if not r: return
        soup = BeautifulSoup(r.text,"html.parser")
        missing = [tag.get("src") or tag.get("href") for tag in soup.find_all(["script","link"])
                   if (tag.get("src") or tag.get("href","")).startswith("http") and not tag.get("integrity")]
        if len(missing) >= 2:
            self._add("Subresource Integrity Missing",f"{len(missing)} external resource(s) lack SRI hashes.",
                "External scripts/styles without SRI hashes — CDN compromise can inject malicious code.")

    def check_open_redirect(self):
        payloads = ["?next=https://evil.com","?redirect=https://evil.com","?url=https://evil.com",
                    "?redir=https://evil.com","?return_url=https://evil.com","?goto=https://evil.com"]
        for p in payloads:
            url = self.target + p
            r = safe_get(url,allow_redirects=False,timeout=6)
            if r and r.status_code in (301,302,307,308):
                loc = r.headers.get("Location","")
                if "evil.com" in loc:
                    self._add("Open Redirect",f"Redirect to {loc} via {url}",
                        "Unvalidated redirect — phishing via trusted domain URL.",[url])
                    return

    def check_xss_indicators(self):
        r = self._main()
        if not r: return
        soup = BeautifulSoup(r.text,"html.parser")
        marker = "xss7z9qtest"
        for form in soup.find_all("form")[:3]:
            action = form.get("action") or self.target
            if not action.startswith("http"):
                action = self.target + "/" + action.lstrip("/")
            inputs = {inp.get("name"):marker for inp in form.find_all("input")
                      if inp.get("name") and inp.get("type","") not in ("submit","hidden","csrf")}
            if not inputs: continue
            try:
                resp = SESSION.post(action,data=inputs,timeout=8,verify=False)
                if marker in resp.text:
                    self._add("No Input Validation (Reflected)",f"Test value reflected from {action}",
                        "User input reflected unencoded — potential Reflected XSS.",[action],severity="High")
                    break
            except Exception:
                pass

    def check_error_disclosure(self):
        error_sigs = {
            "PHP":      ["fatal error","warning:","parse error","stack trace","on line"],
            "Python":   ["traceback","django.core","flask","wsgi","exception in"],
            "Java":     ["javax.","java.lang","at org.","struts","tomcat"],
            "ASP.NET":  ["asp.net","system.web","server error in '/' application"],
            "Database": ["sql syntax","mysql_fetch","ora-","pg::","sqlite","unclosed quotation"],
        }
        for url in [self.target+"/nonexistent-vapt-test",self.target+"/?id=1'"]:
            r = safe_get(url,timeout=6)
            if r:
                body_lower = r.text.lower()
                for tech, sigs in error_sigs.items():
                    if any(s in body_lower for s in sigs):
                        self._add("Error Message Information Leakage",
                            f"Verbose {tech} error at {url}",
                            f"Detailed {tech} errors exposed — reveals internals and potential attack vectors.",[url])
                        return

    def check_jwt(self):
        r = self._main()
        if not r: return
        jwt_pattern = r"eyJ[A-Za-z0-9_-]+\.eyJ[A-Za-z0-9_-]+\.[A-Za-z0-9_-]*"
        all_jwts = re.findall(jwt_pattern,r.text)
        all_jwts += [c.value for c in r.cookies if re.match(jwt_pattern,c.value)]
        for token in all_jwts[:1]:
            try:
                parts = token.split(".")
                header = json.loads(base64.b64decode(parts[0] + "=="))
                alg = header.get("alg","")
                if alg.lower() in ("none",""):
                    self._add("JWT Misconfiguration",f"JWT alg=none found",
                        "JWT 'none' algorithm allows forging tokens without a signature.",severity="Critical")
                elif alg == "HS256":
                    self._add("JWT Misconfiguration",f"JWT using HS256 (symmetric) found",
                        "HS256 tokens vulnerable to brute-force if secret is weak.",severity="Medium")
            except Exception:
                pass

    def check_websocket(self):
        r = self._main()
        if not r: return
        ws = re.findall(r"ws://[^\s\"']+",r.text)
        if ws:
            self._add("WebSocket Without TLS",f"Insecure WebSocket: {ws[0]}",
                "ws:// WebSocket transmits data unencrypted — use wss:// instead.")

    def check_path_traversal(self):
        for p in ["?file=../../../../etc/passwd","?path=../../../etc/shadow","?page=....//....//etc/passwd"]:
            url = self.target + p
            r = safe_get(url,timeout=6)
            if r and ("root:" in r.text or "bin/bash" in r.text):
                self._add("Path Traversal Indicators",f"Traversal response at {url}",
                    "Directory traversal possible — OS files may be readable.",[url],severity="Critical")
                return

    def check_ssrf(self):
        for p in ["?url=http://169.254.169.254/latest/meta-data/","?webhook=http://169.254.169.254/"]:
            url = self.target + p
            r = safe_get(url,timeout=6)
            if r and any(k in r.text for k in ["ami-id","instance-id","security-credentials"]):
                self._add("SSRF Surface Detected",f"SSRF at {url} returned AWS metadata",
                    "Server fetches internal cloud metadata — SSRF confirmed.",[url],severity="Critical")
                return

    def check_default_credentials(self):
        admin_paths = ["/wp-login.php","/admin/login","/login","/admin"]
        default_creds = [("admin","admin"),("admin","password"),("admin","123456")]
        for path in admin_paths[:2]:
            url = self.target + path
            r = safe_get(url,timeout=5)
            if not r or r.status_code != 200: continue
            soup = BeautifulSoup(r.text,"html.parser")
            for form in soup.find_all("form")[:1]:
                action = form.get("action") or url
                if not action.startswith("http"):
                    action = self.target + "/" + action.lstrip("/")
                user_field = next((i.get("name") for i in form.find_all("input")
                                   if i.get("type","") in ("text","email") or "user" in (i.get("name") or "").lower()),None)
                pass_field = next((i.get("name") for i in form.find_all("input") if i.get("type")=="password"),None)
                if user_field and pass_field:
                    for user, pwd in default_creds[:2]:
                        try:
                            resp = SESSION.post(action,data={user_field:user,pass_field:pwd},
                                               timeout=8,verify=False,allow_redirects=True)
                            if resp.status_code == 200 and any(k in resp.text.lower() for k in
                               ["dashboard","logout","welcome","settings"]) and \
                               "invalid" not in resp.text.lower():
                                self._add("Default Credentials Hint",
                                    f"Login at {action} accepted {user}/{pwd}",
                                    "Default credentials accepted — immediate password change required.",[action],severity="Critical")
                                return
                        except Exception:
                            pass

    def check_dns_zone_transfer(self):
        try:
            ns_records = dns.resolver.resolve(self.domain,"NS",lifetime=5)
            for ns in ns_records:
                try:
                    ns_ip = str(dns.resolver.resolve(str(ns),"A",lifetime=4)[0])
                    zone = dns.zone.from_xfr(dns.query.xfr(ns_ip,self.domain,timeout=5))
                    if zone:
                        self._add("DNS Zone Transfer Possible",
                            f"AXFR succeeded from NS {ns} ({ns_ip})",
                            "Full DNS zone exposed — all subdomains and infrastructure revealed.",severity="High")
                        return
                except Exception:
                    pass
        except Exception:
            pass

    def check_dns_email(self):
        try:
            spf_found = False
            try:
                answers = dns.resolver.resolve(self.domain,"TXT",lifetime=5)
                spf_found = any("v=spf1" in str(r) for r in answers)
            except Exception:
                pass
            if not spf_found:
                self._add("SPF Record Missing",f"No SPF TXT for {self.domain}",
                    "Without SPF, anyone can spoof email from your domain.")
            try:
                answers = dns.resolver.resolve(f"_dmarc.{self.domain}","TXT",lifetime=5)
                dmarc_found = any("v=dmarc1" in str(r).lower() for r in answers)
                if not dmarc_found:
                    self._add("DMARC Record Missing",f"No valid DMARC at _dmarc.{self.domain}",
                        "Without DMARC, email spoofing is trivial.")
            except Exception:
                self._add("DMARC Record Missing",f"_dmarc.{self.domain} returned no results",
                    "DMARC policy absent — domain vulnerable to email spoofing.")
            try:
                dns.resolver.resolve(self.domain,"DNSKEY",lifetime=5)
            except (dns.resolver.NoAnswer,dns.resolver.NXDOMAIN):
                self._add("DNSSEC Not Implemented",f"No DNSKEY for {self.domain}",
                    "DNSSEC not enabled — DNS cache poisoning risk.")
            except Exception:
                pass
        except Exception:
            pass

    # ══════════════════════════════════════════════════════════════════════
    #   MASTER RUN
    # ══════════════════════════════════════════════════════════════════════

    def run(self):
        all_steps = [
            ("DNS records",           self.recon_dns),
            ("IP geolocation",        self.recon_ip_info),
            ("WHOIS registration",    self.recon_whois),
            ("SSL certificate",       self.recon_ssl_certificate),
            ("Subdomains (50+)",      self.recon_subdomains),
            ("Port scan (25 ports)",  self.recon_port_scan),
            ("Technology stack",      self.recon_technologies),
            ("WAF detection",         self.recon_waf_detection),
            ("Email harvesting",      self.recon_email_harvest),
            ("Source code comments",  self.recon_source_comments),
            ("Linked pages",          self.recon_linked_pages),
            ("CT log query",          self.recon_ct_logs),
            ("Security headers",      self.check_security_headers),
            ("CORS policy",           self.check_cors),
            ("CORS preflight",        self.check_cors_preflight),
            ("Server disclosure",     self.check_server_disclosure),
            ("SSL/TLS redirect",      self.check_ssl),
            ("Cookies",               self.check_cookies),
            ("Cache control",         self.check_cache_control),
            ("Clickjacking",          self.check_clickjacking),
            ("Mixed content",         self.check_mixed_content),
            ("Directory listing",     self.check_directory_listing),
            ("Sensitive files",       self.check_sensitive_files),
            ("Admin panels",          self.check_admin_panels),
            ("Backup files",          self.check_backup_files),
            ("API endpoints",         self.check_api_exposure),
            ("HTTP methods",          self.check_http_methods),
            ("robots.txt",            self.check_robots),
            ("JS libraries",          self.check_js_libraries),
            ("Subresource integrity", self.check_sri),
            ("Open redirect",         self.check_open_redirect),
            ("Input validation/XSS",  self.check_xss_indicators),
            ("Error disclosure",      self.check_error_disclosure),
            ("JWT tokens",            self.check_jwt),
            ("WebSocket security",    self.check_websocket),
            ("Path traversal",        self.check_path_traversal),
            ("SSRF surface",          self.check_ssrf),
            ("Default credentials",   self.check_default_credentials),
            ("DNS zone transfer",     self.check_dns_zone_transfer),
            ("DNS/email security",    self.check_dns_email),
        ]

        print(Fore.CYAN + f"\n  [*] Target : {self.target}")
        print(Fore.CYAN +  f"  [*] Domain : {self.domain}\n")

        for label, fn in tqdm(all_steps, desc="  Scanning",
                              bar_format="{l_bar}"+Fore.GREEN+"{bar}"+Style.RESET_ALL+"{r_bar}"):
            try:
                fn()
            except Exception:
                pass

        return self.findings, self.intel


# ─────────────────────────────────────────────────────────────────────────────
#  CHART GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def generate_charts(findings, intel, out_dir):
    os.makedirs(out_dir,exist_ok=True)
    charts = {}

    counts = {s:0 for s in SEVERITY_ORDER}
    for f in findings:
        counts[f["severity"]] += 1
    nz = {k:v for k,v in counts.items() if v > 0}

    # Bar chart
    if nz:
        fig,ax = plt.subplots(figsize=(9,4.5))
        bars = ax.bar(list(nz.keys()),list(nz.values()),
                      color=[SEVERITY_COLORS[k] for k in nz],edgecolor="white",linewidth=1,width=0.5)
        ax.set_title("Vulnerability Count by Severity",fontsize=14,fontweight="bold",pad=14)
        ax.set_xlabel("Severity",fontsize=11); ax.set_ylabel("Count",fontsize=11)
        ax.set_facecolor("#F5F5F5"); fig.patch.set_facecolor("#FFFFFF")
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        ax.yaxis.set_major_locator(ticker.MaxNLocator(integer=True))
        for bar,val in zip(bars,nz.values()):
            ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.05,str(val),
                    ha="center",va="bottom",fontweight="bold",fontsize=12)
        plt.tight_layout()
        p = os.path.join(out_dir,"bar_severity.png")
        plt.savefig(p,dpi=150,bbox_inches="tight"); plt.close()
        charts["bar_severity"] = p

    # Pie chart
    if nz:
        fig,ax = plt.subplots(figsize=(6,5))
        ax.pie(list(nz.values()),labels=list(nz.keys()),
               colors=[SEVERITY_COLORS[k] for k in nz],
               autopct="%1.0f%%",startangle=140,
               wedgeprops=dict(edgecolor="white",linewidth=1.5),
               textprops=dict(fontsize=10))
        ax.set_title("Severity Distribution",fontsize=13,fontweight="bold",pad=12)
        plt.tight_layout()
        p = os.path.join(out_dir,"pie_severity.png")
        plt.savefig(p,dpi=150,bbox_inches="tight"); plt.close()
        charts["pie_severity"] = p

    # OWASP breakdown
    owasp_counts = {}
    for f in findings:
        key = f"{f['owasp_id']} {f['owasp_name']}"
        owasp_counts[key] = owasp_counts.get(key,0) + 1
    if owasp_counts:
        sorted_owasp = dict(sorted(owasp_counts.items(),key=lambda x:-x[1]))
        labels = [k[:35] for k in sorted_owasp.keys()]
        vals   = list(sorted_owasp.values())
        fig,ax = plt.subplots(figsize=(11,max(4,len(labels)*0.55+1.5)))
        colors = plt.cm.RdYlGn_r(np.linspace(0.1,0.9,len(labels)))
        bars = ax.barh(labels,vals,color=colors,edgecolor="white")
        ax.set_title("Findings by OWASP Top 10 Category",fontsize=13,fontweight="bold",pad=12)
        ax.set_xlabel("Count",fontsize=11); ax.invert_yaxis()
        ax.set_facecolor("#F5F5F5"); fig.patch.set_facecolor("#FFFFFF")
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        ax.xaxis.set_major_locator(ticker.MaxNLocator(integer=True))
        for bar,val in zip(bars,vals):
            ax.text(bar.get_width()+0.05,bar.get_y()+bar.get_height()/2,str(val),va="center",fontweight="bold")
        plt.tight_layout()
        p = os.path.join(out_dir,"owasp_breakdown.png")
        plt.savefig(p,dpi=150,bbox_inches="tight"); plt.close()
        charts["owasp"] = p

    # Open ports chart
    ports = intel.get("ports",[])
    if ports:
        port_labels = [f"{p['port']}\n{p['service']}" for p in ports]
        fig,ax = plt.subplots(figsize=(max(6,len(ports)*0.9),3.5))
        ax.bar(port_labels,[1]*len(ports),color="#FF6B6B",edgecolor="white")
        ax.set_title("Open Ports Detected",fontsize=13,fontweight="bold")
        ax.set_yticks([]); ax.set_facecolor("#F5F5F5"); fig.patch.set_facecolor("#FFFFFF")
        ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
        plt.tight_layout()
        p = os.path.join(out_dir,"ports.png")
        plt.savefig(p,dpi=150,bbox_inches="tight"); plt.close()
        charts["ports"] = p

    return charts


# ─────────────────────────────────────────────────────────────────────────────
#  DOCX REPORT
# ─────────────────────────────────────────────────────────────────────────────

def styled_p(doc,text,bold=False,size=11,color=None,align=None,space_after=6,italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if align: p.alignment = align
    run = p.add_run(text)
    run.bold=bold; run.italic=italic; run.font.size=Pt(size)
    if color:
        r,g,b = hex_rgb(color)
        run.font.color.rgb = RGBColor(r,g,b)
    return p

def section_heading(doc,text,level=1):
    colors = {1:"1F3864",2:"2E75B6",3:"404040"}
    sizes  = {1:15,2:12,3:10}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text); run.bold=True; run.font.size=Pt(sizes[level])
    r,g,b = hex_rgb(colors[level]); run.font.color.rgb=RGBColor(r,g,b)
    if level <= 2:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6")
        bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),colors[level])
        pBdr.append(bot); pPr.append(pBdr)
    return p

def severity_badge(para,sev):
    fill=SEVERITY_FILL.get(sev,"888888"); txt=SEVERITY_TEXT.get(sev,"FFFFFF")
    run = para.add_run(f"  {sev.upper()}  ")
    run.bold=True; run.font.size=Pt(9)
    r,g,b=hex_rgb(txt); run.font.color.rgb=RGBColor(r,g,b)
    rPr=run._r.get_or_add_rPr()
    shd=OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),fill)
    rPr.append(shd)

def add_kv_table(doc,rows_data):
    tbl=doc.add_table(rows=0,cols=2); tbl.style="Table Grid"
    for key,val in rows_data:
        row=tbl.add_row().cells
        set_cell_bg(row[0],"EBF3FB")
        k=row[0].paragraphs[0].add_run(str(key)); k.bold=True; k.font.size=Pt(9)
        v=row[1].paragraphs[0].add_run(str(val)[:300]); v.font.size=Pt(9)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

def build_report(target,domain,findings,intel,charts,out_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin=Cm(2); section.bottom_margin=Cm(2)
        section.left_margin=Cm(2.5); section.right_margin=Cm(2.5)

    counts = {s:0 for s in SEVERITY_ORDER}
    for f in findings: counts[f["severity"]] += 1
    sorted_findings = sorted(findings,key=lambda x:SEVERITY_ORDER.index(x["severity"]))

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(50)
    run=p.add_run("VAPT ANALYTICAL REPORT"); run.bold=True; run.font.size=Pt(30)
    run.font.color.rgb=RGBColor(0x1F,0x38,0x64)
    styled_p(doc,"Vulnerability Assessment + OSINT Intelligence — Advanced Edition v2.0",
             size=13,color="2E75B6",align=WD_ALIGN_PARAGRAPH.CENTER)
    styled_p(doc,f"Target: {target}",bold=True,size=12,align=WD_ALIGN_PARAGRAPH.CENTER)
    styled_p(doc,f"Assessment Date: {datetime.date.today().strftime('%d %B %Y')}",
             size=11,align=WD_ALIGN_PARAGRAPH.CENTER,space_after=30)
    tbl=doc.add_table(rows=2,cols=5); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,sev in enumerate(SEVERITY_ORDER):
        set_cell_bg(tbl.rows[0].cells[i],SEVERITY_FILL[sev])
        hp=tbl.rows[0].cells[i].paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        hr=hp.add_run(sev); hr.bold=True; hr.font.size=Pt(9)
        r2,g2,b2=hex_rgb(SEVERITY_TEXT[sev]); hr.font.color.rgb=RGBColor(r2,g2,b2)
        vp=tbl.rows[1].cells[i].paragraphs[0]; vp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        vr=vp.add_run(str(counts[sev])); vr.bold=True; vr.font.size=Pt(16)
    doc.add_page_break()

    # ── SECTION 1: TARGET INTELLIGENCE ───────────────────────────────────────
    section_heading(doc,"1. Target Intelligence & OSINT")

    section_heading(doc,"  1.1  IP Address & Hosting Information",level=2)
    ip_rows = []
    for info in intel.get("ip_info",[]):
        ip_rows += [("IP Address",info.get("query",info.get("ip","N/A"))),
                    ("Country/Region/City",f"{info.get('country','N/A')} / {info.get('regionName','N/A')} / {info.get('city','N/A')}"),
                    ("ISP",info.get("isp","N/A")),("Organisation",info.get("org","N/A")),("ASN",info.get("as","N/A"))]
    add_kv_table(doc,ip_rows or [("IP Info","Could not retrieve")])

    section_heading(doc,"  1.2  WHOIS Registration Data",level=2)
    w=intel.get("whois",{})
    add_kv_table(doc,[
        ("Registrar",w.get("registrar","N/A")),("Registered",w.get("creation_date","N/A")),
        ("Expires",w.get("expiry_date","N/A")),("Last Updated",w.get("updated_date","N/A")),
        ("Organisation",w.get("org","N/A")),("Country",w.get("country","N/A")),
        ("Status",w.get("status","N/A")),("Name Servers",", ".join(w.get("name_servers",[])) or "N/A"),
        ("WHOIS Emails",", ".join(w.get("emails",[])) or "None found"),
    ])

    section_heading(doc,"  1.3  SSL/TLS Certificate Details",level=2)
    ssl_d=intel.get("ssl",{})
    subj=ssl_d.get("subject",{}); issuer=ssl_d.get("issuer",{})
    add_kv_table(doc,[
        ("Subject CN",subj.get("CN","N/A")),("Subject Org",subj.get("O","N/A")),
        ("Issuer",f"{issuer.get('CN','N/A')} / {issuer.get('O','N/A')}"),
        ("Protocol",ssl_d.get("protocol","N/A")),("Cipher Suite",str(ssl_d.get("cipher","N/A"))),
        ("Valid Until",ssl_d.get("not_after","N/A")),("Days Until Expiry",str(ssl_d.get("days_until_expiry","N/A"))),
        ("SHA-256 Fingerprint",ssl_d.get("sha256_fp","N/A")[:60]),
        ("SANs","; ".join(ssl_d.get("san",[]))[:200] or "None"),
    ])

    section_heading(doc,"  1.4  Technology Stack Fingerprint",level=2)
    tech=intel.get("technologies",{})
    tech_rows=[(k,str(v) if not isinstance(v,list) else ", ".join(v)) for k,v in tech.items()]
    add_kv_table(doc,tech_rows or [("Technology","Could not fingerprint")])

    waf=intel.get("waf",{})
    if waf:
        section_heading(doc,"  1.5  WAF / Firewall Detection",level=2)
        add_kv_table(doc,[("WAF Detected","Yes" if waf.get("detected") else "No"),
                         ("WAF Product",waf.get("product","None"))])

    section_heading(doc,"  1.6  DNS Records",level=2)
    dns_d=intel.get("dns",{})
    dns_rows=[]
    for rtype in ["A","AAAA","MX","NS","TXT","CNAME","SOA","CAA"]:
        vals=dns_d.get(rtype,[])
        if vals: dns_rows.append((rtype,"\n".join(str(v)[:100] for v in vals[:5])))
    for ip,ptr in list(dns_d.get("PTR_MAP",{}).items())[:3]:
        dns_rows.append((f"PTR ({ip})",ptr))
    add_kv_table(doc,dns_rows or [("DNS","No records resolved")])

    section_heading(doc,"  1.7  Discovered Subdomains",level=2)
    subs=intel.get("subdomains",[])
    if subs:
        tbl2=doc.add_table(rows=1,cols=2); tbl2.style="Table Grid"
        for cell,txt in zip(tbl2.rows[0].cells,["Subdomain","IP Address(es)"]):
            set_cell_bg(cell,"1F3864")
            r2=cell.paragraphs[0].add_run(txt); r2.bold=True; r2.font.size=Pt(9)
            r2.font.color.rgb=RGBColor(255,255,255)
        for s in subs[:30]:
            row=tbl2.add_row().cells
            row[0].paragraphs[0].add_run(s["subdomain"]).font.size=Pt(9)
            row[1].paragraphs[0].add_run(", ".join(s["ips"])).font.size=Pt(9)
    else:
        styled_p(doc,"No active subdomains discovered.",size=10,italic=True)
    doc.add_paragraph()

    section_heading(doc,"  1.8  Open Ports",level=2)
    ports=intel.get("ports",[])
    if ports:
        tbl3=doc.add_table(rows=1,cols=2); tbl3.style="Table Grid"
        for cell,txt in zip(tbl3.rows[0].cells,["Port","Service"]):
            set_cell_bg(cell,"1F3864")
            r3=cell.paragraphs[0].add_run(txt); r3.bold=True; r3.font.size=Pt(9)
            r3.font.color.rgb=RGBColor(255,255,255)
        for p in ports:
            row=tbl3.add_row().cells
            row[0].paragraphs[0].add_run(str(p["port"])).font.size=Pt(9)
            row[1].paragraphs[0].add_run(p["service"]).font.size=Pt(9)
    else:
        styled_p(doc,"No unusual open ports detected.",size=10,italic=True)
    doc.add_paragraph()

    section_heading(doc,"  1.9  Certificate Transparency Logs (crt.sh)",level=2)
    ct=intel.get("ct_logs",[])
    styled_p(doc,f"{len(ct)} certificate entry(ies) found in CT logs.",size=10)
    if ct:
        tbl4=doc.add_table(rows=1,cols=2); tbl4.style="Table Grid"
        for cell,txt in zip(tbl4.rows[0].cells,["Domain / SAN","Date Logged"]):
            set_cell_bg(cell,"1F3864")
            r4=cell.paragraphs[0].add_run(txt); r4.bold=True; r4.font.size=Pt(9)
            r4.font.color.rgb=RGBColor(255,255,255)
        for c in ct[:20]:
            row=tbl4.add_row().cells
            row[0].paragraphs[0].add_run(c["domain"][:80]).font.size=Pt(9)
            row[1].paragraphs[0].add_run(c["logged"]).font.size=Pt(9)
    doc.add_paragraph()

    emails=intel.get("emails",[])
    if emails:
        section_heading(doc,"  1.10  Harvested Email Addresses",level=2)
        for e in emails[:20]:
            bp=doc.add_paragraph(style="List Bullet"); bp.add_run(e).font.size=Pt(10)

    comments=intel.get("source_comments",[])
    if comments:
        section_heading(doc,"  1.11  Sensitive HTML Comments",level=2)
        for c in comments[:5]:
            bp=doc.add_paragraph(style="List Bullet"); bp.add_run(c[:200]).font.size=Pt(9)

    doc.add_page_break()

    # ── SECTION 2: EXECUTIVE SUMMARY ─────────────────────────────────────────
    section_heading(doc,"2. Executive Summary")
    total=len(findings); crit=counts["Critical"]; high=counts["High"]
    med=counts["Medium"]; low=counts["Low"]
    summary=(f"An automated VAPT assessment of {target} identified {total} security finding(s): "
             f"{crit} Critical, {high} High, {med} Medium, and {low} Low severity. ")
    if crit > 0: summary += "CRITICAL issues require immediate emergency remediation. "
    if high > 0: summary += "High severity findings should be addressed within 7 days. "
    if med > 0:  summary += "Medium issues should be resolved within 30 days. "
    summary += ("Assessment covered security headers, SSL/TLS, cookie security, CORS, injection testing, "
                "authentication controls, open port scanning, DNS/email security, subdomain enumeration, "
                "WAF detection, CT log analysis, and full OSINT intelligence gathering.")
    styled_p(doc,summary,size=11,space_after=8)

    # ── SECTION 3: CHARTS ─────────────────────────────────────────────────────
    doc.add_page_break()
    section_heading(doc,"3. Vulnerability Distribution & Analysis")
    for key in ["bar_severity","pie_severity"]:
        if key in charts and os.path.exists(charts[key]):
            doc.add_picture(charts[key],width=Inches(5.2))
            doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if "owasp" in charts and os.path.exists(charts["owasp"]):
        section_heading(doc,"   OWASP Top 10 Breakdown",level=2)
        doc.add_picture(charts["owasp"],width=Inches(6.0))
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    if "ports" in charts and os.path.exists(charts["ports"]):
        section_heading(doc,"   Open Ports",level=2)
        doc.add_picture(charts["ports"],width=Inches(5.0))
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER

    # ── SECTION 4: SUMMARY TABLE ──────────────────────────────────────────────
    doc.add_page_break()
    section_heading(doc,"4. Summary of Vulnerabilities")
    tbl5=doc.add_table(rows=1,cols=4); tbl5.style="Table Grid"
    for cell,txt in zip(tbl5.rows[0].cells,["ID","Vulnerability","Severity","OWASP"]):
        set_cell_bg(cell,"1F3864"); p5=cell.paragraphs[0]; p5.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r5=p5.add_run(txt); r5.bold=True; r5.font.size=Pt(9); r5.font.color.rgb=RGBColor(255,255,255)
    for i,f in enumerate(sorted_findings):
        row=tbl5.add_row().cells; fill="F2F2F2" if i%2 else "FFFFFF"
        for cell in row: set_cell_bg(cell,fill)
        row[0].paragraphs[0].add_run(f["id"]).font.size=Pt(9)
        row[1].paragraphs[0].add_run(f["name"]).font.size=Pt(9)
        sp=row[2].paragraphs[0]; sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(row[2],SEVERITY_FILL[f["severity"]])
        sr=sp.add_run(f["severity"]); sr.bold=True; sr.font.size=Pt(9)
        r6,g6,b6=hex_rgb(SEVERITY_TEXT[f["severity"]]); sr.font.color.rgb=RGBColor(r6,g6,b6)
        row[3].paragraphs[0].add_run(f"{f['owasp_id']} – {f['owasp_name']}").font.size=Pt(9)
    doc.add_paragraph()

    # ── SECTION 5: DETAILED FINDINGS ─────────────────────────────────────────
    doc.add_page_break()
    section_heading(doc,"5. Detailed Vulnerability Findings")
    for f in sorted_findings:
        section_heading(doc,f"  {f['id']}  {f['name']}",level=2)
        bp=doc.add_paragraph(); bp.paragraph_format.space_after=Pt(3)
        bp.add_run("Severity: ").font.size=Pt(10); severity_badge(bp,f["severity"])
        op=doc.add_paragraph(); op.paragraph_format.space_after=Pt(8)
        or2=op.add_run(f"OWASP {f['owasp_id']} – {f['owasp_name']}"); or2.italic=True; or2.font.size=Pt(10)
        r7,g7,b7=hex_rgb("2E75B6"); or2.font.color.rgb=RGBColor(r7,g7,b7)
        for label,content in [("Summary",f["detail"]),("Evidence",f["evidence"]),("Remediation",f["remedy"])]:
            lp=doc.add_paragraph(); lp.paragraph_format.space_after=Pt(2); lp.paragraph_format.space_before=Pt(4)
            lr=lp.add_run(label+":"); lr.bold=True; lr.font.size=Pt(10)
            r8,g8,b8=hex_rgb("1F3864"); lr.font.color.rgb=RGBColor(r8,g8,b8)
            body=doc.add_paragraph(); body.paragraph_format.space_after=Pt(4); body.paragraph_format.left_indent=Pt(14)
            body.add_run(content).font.size=Pt(10)
        up=doc.add_paragraph(); up.paragraph_format.left_indent=Pt(14); up.paragraph_format.space_after=Pt(14)
        ur=up.add_run("Affected URL(s): "+", ".join(f["urls"])); ur.font.size=Pt(9)
        r9,g9,b9=hex_rgb("0000CC"); ur.font.color.rgb=RGBColor(r9,g9,b9)

    # ── SECTION 6: REMEDIATION ROADMAP ───────────────────────────────────────
    doc.add_page_break()
    section_heading(doc,"6. Prioritised Remediation Roadmap")
    priority_desc={"Critical":"Address immediately — within 24 hours","High":"Address urgently — within 7 days",
                   "Medium":"Address in next sprint — within 30 days","Low":"Schedule for next quarter","Info":"Informational"}
    for sev in SEVERITY_ORDER:
        sev_f=[f for f in findings if f["severity"]==sev]
        if not sev_f: continue
        section_heading(doc,f"  {sev}  —  {priority_desc[sev]}",level=2)
        for f in sev_f:
            bp=doc.add_paragraph(style="List Bullet"); bp.paragraph_format.space_after=Pt(3)
            br=bp.add_run(f"{f['name']}: "); br.bold=True; br.font.size=Pt(10)
            bp.add_run(f["remedy"]).font.size=Pt(10)

    # ── SECTION 7: APPENDIX ───────────────────────────────────────────────────
    doc.add_page_break()
    section_heading(doc,"7. Appendix — Internal Links Discovered")
    for link in intel.get("internal_links",[])[:30]:
        bp=doc.add_paragraph(style="List Bullet"); bp.add_run(link[:120]).font.size=Pt(9)

    # Footer
    for section in doc.sections:
        fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run(f"VAPT Analytical Report v2.0  |  {domain}  |  {datetime.date.today().strftime('%d %B %Y')}  |  CONFIDENTIAL"
                  ).font.size=Pt(8)

    doc.save(out_path)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
#  CONSOLE SUMMARY
# ─────────────────────────────────────────────────────────────────────────────

def print_summary(findings,intel):
    COLOR={"Critical":Fore.RED,"High":Fore.LIGHTYELLOW_EX,"Medium":Fore.YELLOW,"Low":Fore.WHITE,"Info":Fore.CYAN}
    print(Fore.CYAN+"\n"+"═"*65)
    print(Fore.CYAN+"  OSINT INTELLIGENCE SUMMARY")
    print(Fore.CYAN+"═"*65)
    for ip_info in intel.get("ip_info",[]):
        print(f"  {Fore.GREEN}IP{Style.RESET_ALL}          {ip_info.get('query',ip_info.get('ip','?'))}")
        print(f"  {Fore.GREEN}Location{Style.RESET_ALL}    {ip_info.get('city','?')}, {ip_info.get('regionName','?')}, {ip_info.get('country','?')}")
        print(f"  {Fore.GREEN}ISP/ASN{Style.RESET_ALL}     {ip_info.get('isp','?')} | {ip_info.get('as','?')}")
    tech=intel.get("technologies",{})
    for key in ["CMS","WAF","Server","JS Frameworks","CDN","Email Provider"]:
        if tech.get(key):
            print(f"  {Fore.GREEN}{key:<12}{Style.RESET_ALL}  {tech[key]}")
    subs=intel.get("subdomains",[])
    print(f"  {Fore.GREEN}Subdomains{Style.RESET_ALL}   {len(subs)} discovered")
    ports=intel.get("ports",[])
    if ports:
        port_str=", ".join([f"{p['port']}/{p['service']}" for p in ports])
        print(f"  {Fore.GREEN}Open Ports{Style.RESET_ALL}   {port_str}")
    ssl_d=intel.get("ssl",{})
    if ssl_d.get("days_until_expiry") is not None:
        days=ssl_d["days_until_expiry"]
        col=Fore.RED if days<30 else Fore.GREEN
        print(f"  {Fore.GREEN}SSL Expiry{Style.RESET_ALL}   {col}{days} days{Style.RESET_ALL}")
    emails=intel.get("emails",[])
    if emails:
        print(f"  {Fore.GREEN}Emails{Style.RESET_ALL}       {', '.join(emails[:3])}")
    ct=intel.get("ct_logs",[])
    print(f"  {Fore.GREEN}CT Log Entries{Style.RESET_ALL} {len(ct)}")
    print(Fore.CYAN+"\n"+"═"*65)
    print(Fore.CYAN+f"  VULNERABILITY FINDINGS ({len(findings)} total)")
    print(Fore.CYAN+"═"*65)
    for f in sorted(findings,key=lambda x:SEVERITY_ORDER.index(x["severity"])):
        c=COLOR.get(f["severity"],Fore.WHITE)
        print(f"  {c}[{f['severity']:<8}]{Style.RESET_ALL}  {f['id']}  {f['name']}")
    print(Fore.CYAN+"─"*65)
    counts={s:0 for s in SEVERITY_ORDER}
    for f in findings: counts[f["severity"]] += 1
    for sev in SEVERITY_ORDER:
        if counts[sev]: print(f"  {COLOR[sev]}{sev:<12}{Style.RESET_ALL}  {counts[sev]}")
    print()


# ─────────────────────────────────────────────────────────────────────────────
#  ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

def main():
    banner()
    if len(sys.argv) < 2:
        print(Fore.RED+"  Usage: python vapt_scanner.py <target_url>")
        print(Fore.RED+"  Example: python vapt_scanner.py https://example.com\n")
        sys.exit(1)

    target = sys.argv[1].strip()
    domain = get_domain(normalise_url(target))

    scanner = VAPTScanner(target)
    findings, intel = scanner.run()

    print_summary(findings, intel)

    out_dir = f"vapt_report_{domain}_{datetime.date.today().isoformat()}"
    os.makedirs(out_dir, exist_ok=True)

    print(Fore.CYAN + "[*] Generating charts …")
    charts = generate_charts(findings, intel, out_dir)

    report_path = os.path.join(out_dir, f"VAPT_Report_{domain}.docx")
    print(Fore.CYAN + "[*] Building DOCX report …")
    build_report(normalise_url(target), domain, findings, intel, charts, report_path)

    print(Fore.GREEN + f"\n✅  Report saved  →  {report_path}")
    print(Fore.GREEN + f"    Charts saved   →  {out_dir}/\n")


if __name__ == "__main__":
    main()
